from http.server import BaseHTTPRequestHandler
import json, io, base64
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top','left','bottom','right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), '000000')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def cell_text(cell, text, bold=False, italic=False, underline=False, align='center', size=9):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if align == 'center' else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text) if text else '')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.underline = underline
    run.font.name = '맑은 고딕'
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_gray_cell(cell, text, size=9):
    set_cell_bg(cell, 'D9D9D9')
    cell_text(cell, text, bold=True, size=size)

def make_header_table(doc, d, doc_label, page_str):
    # 헤더 테이블: 문서번호 | 문서유형 | 작성 | | 검토 | | 승인 |
    tbl = doc.add_table(rows=3, cols=8)
    tbl.style = 'Table Grid'
    widths = [1800, 3600, 600, 1200, 600, 1200, 600, 1200]
    
    # Row 0
    r0 = tbl.rows[0]
    add_gray_cell(r0.cells[0], '문 서 번 호', 8)
    add_gray_cell(r0.cells[1], doc_label, 8)
    add_gray_cell(r0.cells[2], '작\n성', 8)
    cell_text(r0.cells[3], '', size=8)
    add_gray_cell(r0.cells[4], '검\n토', 8)
    cell_text(r0.cells[5], '', size=8)
    add_gray_cell(r0.cells[6], '승\n인', 8)
    cell_text(r0.cells[7], '', size=8)
    
    # Row 1
    r1 = tbl.rows[1]
    cell_text(r1.cells[0], d.get('docNo',''), size=8)
    cell_text(r1.cells[1], d.get('productName',''), bold=True, size=9)
    add_gray_cell(r1.cells[2], '작\n성', 8)
    cell_text(r1.cells[3], f"/ {d.get('drafter','')}", size=8)
    add_gray_cell(r1.cells[4], '검\n토', 8)
    cell_text(r1.cells[5], f"/ {d.get('reviewer','')}", size=8)
    add_gray_cell(r1.cells[6], '승\n인', 8)
    cell_text(r1.cells[7], f"/ {d.get('approver','')}", size=8)
    
    # Row 2
    r2 = tbl.rows[2]
    add_gray_cell(r2.cells[0], '제정일자', 8)
    cell_text(r2.cells[1], d.get('establishDate',''), size=8)
    add_gray_cell(r2.cells[2], '개정일자', 8)
    # merge cells 3+4 for revision date
    r2.cells[3].merge(r2.cells[4])
    cell_text(r2.cells[3], d.get('revisionDate',''), size=8)
    add_gray_cell(r2.cells[5], '개정번호', 8)
    cell_text(r2.cells[6], d.get('revisionNo',''), size=8)
    cell_text(r2.cells[7], page_str, size=8)
    
    return tbl

def generate_docx(d):
    from docx.shared import Cm, Pt
    doc = Document()
    
    # 페이지 설정 A4
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    
    fc = d.get('flowConditions', {})
    doc_label = '제조법 문서 (Be품)' if d.get('docType','Be품') == 'Be품' else '제조법 문서'
    
    materials = d.get('materials', [])
    specs = d.get('specs', [])
    
    # ── PAGE 1: 헤더 + Material Balance + 제품규격 ──
    # 우측 그린케미칼(주)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('그린케미칼(주)')
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    
    make_header_table(doc, d, doc_label, '1/4')
    doc.add_paragraph()
    
    # 1.0 Material Balance
    p = doc.add_paragraph(' 1.0 Material Balance ')
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(9)
    
    mat_tbl = doc.add_table(rows=2+max(len(materials),10)+1, cols=4)
    mat_tbl.style = 'Table Grid'
    
    # 헤더행
    hr = mat_tbl.rows[0]
    add_gray_cell(hr.cells[0], '', 8)
    add_gray_cell(hr.cells[1], '원       료       명', 8)
    add_gray_cell(hr.cells[2], '투   입   량 (kg)', 8)
    add_gray_cell(hr.cells[3], '비      고', 8)
    
    # 원료 행
    for i, mat in enumerate(materials):
        row = mat_tbl.rows[i+1]
        cell_text(row.cells[0], mat.get('order',''), size=9)
        changed = mat.get('changed', False)
        cell_text(row.cells[1], mat.get('name',''), bold=changed, italic=changed, underline=changed, size=9)
        cell_text(row.cells[2], mat.get('amount',''), bold=changed, italic=changed, underline=changed, size=9)
        cell_text(row.cells[3], mat.get('note',''), size=9)
    
    # 빈 행
    for i in range(max(10-len(materials), 0)):
        row = mat_tbl.rows[len(materials)+1+i]
        for c in row.cells: cell_text(c, '', size=9)
    
    # TOTAL 행
    total_row = mat_tbl.rows[-1]
    cell_text(total_row.cells[0], '', size=9)
    cell_text(total_row.cells[1], 'TOTAL', bold=True, size=9)
    total_changed = True
    cell_text(total_row.cells[2], d.get('totalAmount',''), bold=total_changed, italic=total_changed, underline=total_changed, size=9)
    cell_text(total_row.cells[3], '', size=9)
    
    doc.add_paragraph()
    
    # 2.0 제품규격
    p = doc.add_paragraph(' 2.0 제품규격')
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(9)
    
    spec_tbl = doc.add_table(rows=1+max(len(specs),8), cols=4)
    spec_tbl.style = 'Table Grid'
    
    sr = spec_tbl.rows[0]
    add_gray_cell(sr.cells[0], '', 8)
    add_gray_cell(sr.cells[1], '항          목', 8)
    add_gray_cell(sr.cells[2], '규          격', 8)
    add_gray_cell(sr.cells[3], '시   험   법', 8)
    
    for i, spec in enumerate(specs):
        row = spec_tbl.rows[i+1]
        cell_text(row.cells[0], spec.get('order',''), size=9)
        cell_text(row.cells[1], spec.get('item',''), size=9)
        changed = spec.get('changed', False)
        cell_text(row.cells[2], spec.get('spec',''), bold=changed, italic=changed, underline=changed, size=9)
        cell_text(row.cells[3], spec.get('method',''), size=9)
    
    for i in range(max(8-len(specs), 0)):
        row = spec_tbl.rows[len(specs)+1+i]
        for c in row.cells: cell_text(c, '', size=9)
    
    doc.add_paragraph()
    
    # 개정/배포처
    rev_tbl = doc.add_table(rows=2, cols=2)
    rev_tbl.style = 'Table Grid'
    add_gray_cell(rev_tbl.rows[0].cells[0], '개  정', 8)
    cell_text(rev_tbl.rows[0].cells[1], d.get('revisionNote',''), align='left', size=9)
    add_gray_cell(rev_tbl.rows[1].cells[0], '배포처', 8)
    cell_text(rev_tbl.rows[1].cells[1], d.get('distributionMfg',''), align='left', size=9)
    
    p = doc.add_paragraph('FB01-06(1)')
    p.runs[0].font.size = Pt(7)
    
    p = doc.add_paragraph('그린케미칼(주)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(10)
    
    # ── PAGE 2: 작업개요 공정 1~5 ──
    doc.add_page_break()
    
    sub_tbl = doc.add_table(rows=2, cols=2)
    sub_tbl.style = 'Table Grid'
    add_gray_cell(sub_tbl.rows[0].cells[0], doc_label, 8)
    cell_text(sub_tbl.rows[0].cells[1], d.get('productName',''), bold=True, size=9)
    add_gray_cell(sub_tbl.rows[1].cells[0], '제정일자', 8)
    cell_text(sub_tbl.rows[1].cells[1], 
        f"{d.get('establishDate','')}  개정일자 {d.get('revisionDate','')}  개정번호 {d.get('revisionNo','')}  페이지 2/4",
        align='left', size=8)
    
    doc.add_paragraph()
    p = doc.add_paragraph(' 3.0 작업개요')
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(9)
    
    flow1_tbl = doc.add_table(rows=2, cols=2)
    flow1_tbl.style = 'Table Grid'
    add_gray_cell(flow1_tbl.rows[0].cells[0], 'FLOW SHEET', 8)
    add_gray_cell(flow1_tbl.rows[0].cells[1], '제    조    조    건', 8)
    
    flow_text = f"{fc.get('starter','Starter')}  {fc.get('catalyst','촉매')}\n↓\n1  원료사입\n↓\n2  치   환\n↓\n3  승   온\n{fc.get('reactant','Reactant')}\n↓\n4  반   응\n↓\n5  숙    성"
    cell = flow1_tbl.rows[1].cells[0]
    cell.text = flow_text
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.name = '맑은 고딕'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cond_text = f"""<준 비> : 사용할 중합조가 충분히 세척, 건조되어 있는가를 확인한다.

1. 원료사입
    지시량의 {fc.get('starter','Starter')}, {fc.get('catalyst','촉매')}를 사입한 후 교반을 실시한다.

2. 치   환
    0 kg/cm²G  ↔  -1 kg/cm²G 로 질소감압치환을 3회 실시한 후 최종압력은 -1.0 kg/cm²G 로 한다

3. 승  온
    반응온도까지 승온한다.({fc.get('reactionTemp','140')}℃)

4. 반   응
    아래의 조건으로 {fc.get('reactant','Reactant')} 반응을 실시한다.
    * 반응온도 : {fc.get('reactionTemp','140')} ± 5℃
    * 반응압력 : {fc.get('reactionPressure','4')} kg/cm²G 이하
    * 반응시간 : 지시량의 {fc.get('reactant','Reactant')} 사입 종료까지

5. 숙   성
    지시량의 {fc.get('reactant','Reactant')} 사입이 종료되면 아래의 조건에서 숙성을 실시한다.
    * 숙성온도 : {fc.get('agingTemp','140')} ± 5℃
    * 숙성압력 : {fc.get('agingPressure','4')} kg/cm²G 이하
    * 숙성시간 : 동일온도에서 압력평형상태가 30분간 지속될 때 까지"""
    
    cond_cell = flow1_tbl.rows[1].cells[1]
    cond_cell.text = cond_text
    for p in cond_cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(8.5)
            run.font.name = '맑은 고딕'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p = doc.add_paragraph('FB01-06(1)')
    p.runs[0].font.size = Pt(7)
    p = doc.add_paragraph('그린케미칼(주)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(10)
    
    # ── PAGE 3: 작업개요 공정 6~9 ──
    doc.add_page_break()
    
    sub_tbl2 = doc.add_table(rows=2, cols=2)
    sub_tbl2.style = 'Table Grid'
    add_gray_cell(sub_tbl2.rows[0].cells[0], doc_label, 8)
    cell_text(sub_tbl2.rows[0].cells[1], d.get('productName',''), bold=True, size=9)
    add_gray_cell(sub_tbl2.rows[1].cells[0], '제정일자', 8)
    cell_text(sub_tbl2.rows[1].cells[1],
        f"{d.get('establishDate','')}  개정일자 {d.get('revisionDate','')}  개정번호 {d.get('revisionNo','')}  페이지 3/4",
        align='left', size=8)
    
    doc.add_paragraph()
    p = doc.add_paragraph(' 3.0 작업개요')
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(9)
    
    flow2_tbl = doc.add_table(rows=2, cols=2)
    flow2_tbl.style = 'Table Grid'
    add_gray_cell(flow2_tbl.rows[0].cells[0], 'FLOW SHEET', 8)
    add_gray_cell(flow2_tbl.rows[0].cells[1], '제    조    조    건', 8)
    
    cell_text(flow2_tbl.rows[1].cells[0], '6  분   석\n↓\n7  냉   각\n↓\n8  탈   취\n↓\n9  포  장', size=8)
    
    spec_lines = '\n'.join([f"       {s.get('item','').strip()}    : {s.get('spec','')}" for s in specs if s.get('item') and s.get('spec')])
    cond2_text = f"""6. 분   석
    숙성이 종료되면 다음 항목을 분석한다.
{spec_lines}

7. 냉   각
    분석치가 규격내인 것을 확인한 후 탈취온도까지 냉각한다.

8. 탈   취
    탈취온도({fc.get('deodorTemp','80')}±5℃)에서 30분간 질소로 탈취한다.

9. 포  장
    본 제품은 Be품이므로 제품의 변질을 방지하기 위하여 질소를 충진하여
    포장한다.(포장온도 : {fc.get('packingTemp','60')}℃ 이하)"""
    
    cond2_cell = flow2_tbl.rows[1].cells[1]
    cond2_cell.text = cond2_text
    for p in cond2_cell.paragraphs:
        for run in p.runs:
            run.font.size = Pt(8.5)
            run.font.name = '맑은 고딕'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p = doc.add_paragraph('FB01-06(1)')
    p.runs[0].font.size = Pt(7)
    p = doc.add_paragraph('그린케미칼(주)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(10)
    
    # ── PAGE 4: 제조기계/유해물질/포장/저장 ──
    doc.add_page_break()
    
    sub_tbl3 = doc.add_table(rows=2, cols=2)
    sub_tbl3.style = 'Table Grid'
    add_gray_cell(sub_tbl3.rows[0].cells[0], doc_label, 8)
    cell_text(sub_tbl3.rows[0].cells[1], d.get('productName',''), bold=True, size=9)
    add_gray_cell(sub_tbl3.rows[1].cells[0], '제정일자', 8)
    cell_text(sub_tbl3.rows[1].cells[1],
        f"{d.get('establishDate','')}  개정일자 {d.get('revisionDate','')}  개정번호 {d.get('revisionNo','')}  페이지 4/4",
        align='left', size=8)
    
    doc.add_paragraph()
    info_tbl = doc.add_table(rows=1, cols=1)
    info_tbl.style = 'Table Grid'
    info_cell = info_tbl.rows[0].cells[0]
    info_cell.text = ''
    lines = [
        ('4.0 제조기계', True),
        (f"    {fc.get('reactorNo','V-302')} Reactor Type", False),
        ('', False),
        ('5.0 유해물질명', True),
        (f"    {fc.get('hazardous','해당없음')}", False),
        ('', False),
        ('6.0 포장용기', True),
        (f"    {fc.get('packageType','Steel Drum (Net.Wt. : 230 kg)')}", False),
        ('', False),
        ('7.0 저장, 보존, 취급, 폐기에 관한 사항', True),
        ('  7.1 저장 및 보존', True),
        (f"      {fc.get('storage','')}", False),
        ('  7.2 취  급', True),
        (f"      {fc.get('handling','')}", False),
        ('  7.3 폐  기', True),
        (f"      {fc.get('disposal','')}", False),
    ]
    for i, (line, bold) in enumerate(lines):
        if i == 0:
            p = info_cell.paragraphs[0]
        else:
            p = info_cell.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9)
        run.font.bold = bold
        run.font.name = '맑은 고딕'
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    p = doc.add_paragraph('FB01-06(1)')
    p.runs[0].font.size = Pt(7)
    p = doc.add_paragraph('그린케미칼(주)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(10)
    
    # ── PAGE 5: 제품보증규격 ──
    doc.add_page_break()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('그린케미칼(주)')
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    
    # 제품보증규격 헤더
    qa_doc_no = (d.get('docNo','') or '').replace('R41','R32')
    q_tbl = doc.add_table(rows=3, cols=8)
    q_tbl.style = 'Table Grid'
    
    r0 = q_tbl.rows[0]
    add_gray_cell(r0.cells[0], '문 서 번 호', 8)
    add_gray_cell(r0.cells[1], '제품규격 (Be품)', 8)
    add_gray_cell(r0.cells[2], '작\n성', 8)
    cell_text(r0.cells[3], '', size=8)
    add_gray_cell(r0.cells[4], '검\n토', 8)
    cell_text(r0.cells[5], '', size=8)
    add_gray_cell(r0.cells[6], '승\n인', 8)
    cell_text(r0.cells[7], '', size=8)
    
    r1 = q_tbl.rows[1]
    cell_text(r1.cells[0], qa_doc_no, size=8)
    cell_text(r1.cells[1], d.get('productName',''), bold=True, size=9)
    add_gray_cell(r1.cells[2], '작\n성', 8)
    cell_text(r1.cells[3], f"/ {d.get('drafter','')}", size=8)
    add_gray_cell(r1.cells[4], '검\n토', 8)
    cell_text(r1.cells[5], f"/ {d.get('reviewer','')}", size=8)
    add_gray_cell(r1.cells[6], '승\n인', 8)
    cell_text(r1.cells[7], f"/ {d.get('approver','')}", size=8)
    
    r2 = q_tbl.rows[2]
    add_gray_cell(r2.cells[0], '제정일자', 8)
    cell_text(r2.cells[1], d.get('establishDate',''), size=8)
    add_gray_cell(r2.cells[2], '개정일자', 8)
    r2.cells[3].merge(r2.cells[4])
    cell_text(r2.cells[3], d.get('revisionDate',''), size=8)
    add_gray_cell(r2.cells[5], '개정번호', 8)
    cell_text(r2.cells[6], d.get('revisionNo',''), size=8)
    cell_text(r2.cells[7], '1/1', size=8)
    
    doc.add_paragraph()
    
    # 공정도 + 분석표
    qa_tbl = doc.add_table(rows=2, cols=6)
    qa_tbl.style = 'Table Grid'
    hr = qa_tbl.rows[0]
    add_gray_cell(hr.cells[0], '공  정  도', 8)
    add_gray_cell(hr.cells[1], '분 석 점', 8)
    add_gray_cell(hr.cells[2], '분 석 항 목', 8)
    add_gray_cell(hr.cells[3], '분 석 법', 8)
    add_gray_cell(hr.cells[4], '규    격', 8)
    add_gray_cell(hr.cells[5], '비    고', 8)
    
    dr = qa_tbl.rows[1]
    flow_str = '1  원료사입\n2  치    환\n3  승    온\n4  반    응\n5  숙    성\n6  분    석\n7  냉    각\n8  탈    취\n9  포    장'
    dr.cells[0].text = flow_str
    for p in dr.cells[0].paragraphs:
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.name = '맑은 고딕'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cell_text(dr.cells[1], '6', bold=True, size=11)
    
    items_text = '\n'.join([s.get('item','') for s in specs if s.get('item')])
    methods_text = '\n'.join([s.get('method','') for s in specs if s.get('method')])
    
    dr.cells[2].text = items_text
    dr.cells[3].text = methods_text
    
    specs_cell = dr.cells[4]
    specs_cell.text = ''
    for s in specs:
        if s.get('spec'):
            p = specs_cell.add_paragraph() if specs_cell.paragraphs[0].text else specs_cell.paragraphs[0]
            run = p.add_run(s['spec'])
            run.font.size = Pt(9)
            run.font.bold = s.get('changed', False)
            run.font.italic = s.get('changed', False)
            run.font.underline = s.get('changed', False)
            run.font.name = '맑은 고딕'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    cell_text(dr.cells[5], '', size=9)
    
    doc.add_paragraph()
    
    rev2_tbl = doc.add_table(rows=2, cols=2)
    rev2_tbl.style = 'Table Grid'
    add_gray_cell(rev2_tbl.rows[0].cells[0], '개  정', 8)
    cell_text(rev2_tbl.rows[0].cells[1], d.get('revisionNote',''), align='left', size=9)
    add_gray_cell(rev2_tbl.rows[1].cells[0], '배포처', 8)
    cell_text(rev2_tbl.rows[1].cells[1], d.get('distributionQuality',''), align='left', size=9)
    
    p = doc.add_paragraph('FB01-07(1)')
    p.runs[0].font.size = Pt(7)
    p = doc.add_paragraph('그린케미칼(주)')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.bold = True
    p.runs[0].font.size = Pt(10)
    
    # ── PAGE 6: 품질보증규격 ──
    doc.add_page_break()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('그린케미칼(주)')
    run.font.bold = True
    run.font.size = Pt(10)
    run.font.name = '맑은 고딕'
    
    qb_tbl = doc.add_table(rows=3, cols=8)
    qb_tbl.style = 'Table Grid'
    
    r0 = qb_tbl.rows[0]
    add_gray_cell(r0.cells[0], '문 서 번 호', 8)
    add_gray_cell(r0.cells[1], '품질보증규격', 8)
    add_gray_cell(r0.cells[2], '작\n성', 8)
    cell_text(r0.cells[3], '', size=8)
    add_gray_cell(r0.cells[4], '검\n토', 8)
    cell_text(r0.cells[5], '', size=8)
    add_gray_cell(r0.cells[6], '승\n인', 8)
    cell_text(r0.cells[7], '', size=8)
    
    r1 = qb_tbl.rows[1]
    cell_text(r1.cells[0], d.get('qaDocNo',''), size=8)
    cell_text(r1.cells[1], f"KONION {d.get('productName','')}", bold=True, size=9)
    add_gray_cell(r1.cells[2], '작\n성', 8)
    cell_text(r1.cells[3], f"/ {d.get('drafter','')}", size=8)
    add_gray_cell(r1.cells[4], '검\n토', 8)
    cell_text(r1.cells[5], f"/ {d.get('reviewer','')}", size=8)
    add_gray_cell(r1.cells[6], '승\n인', 8)
    cell_text(r1.cells[7], f"/ {d.get('approver','')}", size=8)
    
    r2 = qb_tbl.rows[2]
    add_gray_cell(r2.cells[0], '제정일자', 8)
    cell_text(r2.cells[1], d.get('establishDate',''), size=8)
    add_gray_cell(r2.cells[2], '개정일자', 8)
    r2.cells[3].merge(r2.cells[4])
    cell_text(r2.cells[3], d.get('qaRevisionDate','') or d.get('revisionDate',''), size=8)
    add_gray_cell(r2.cells[5], '개정번호', 8)
    cell_text(r2.cells[6], d.get('qaRevisionNo','') or d.get('revisionNo',''), size=8)
    cell_text(r2.cells[7], '1/1', size=8)
    
    doc.add_paragraph()
    
    qa_specs = d.get('qaSpecs', [])
    qc_tbl = doc.add_table(rows=1+len(qa_specs), cols=2)
    qc_tbl.style = 'Table Grid'
    add_gray_cell(qc_tbl.rows[0].cells[0], '구 분 / 분석항목', 8)
    add_gray_cell(qc_tbl.rows[0].cells[1], d.get('productName',''), 8)
    
    for i, qs in enumerate(qa_specs):
        row = qc_tbl.rows[i+1]
        cell_text(row.cells[0], qs.get('item',''), size=9)
        changed = qs.get('changed', False)
        cell_text(row.cells[1], qs.get('spec',''), bold=changed, italic=changed, underline=changed, size=9)
    
    doc.add_paragraph()
    
    sp_tbl = doc.add_table(rows=1, cols=1)
    sp_tbl.style = 'Table Grid'
    sp_cell = sp_tbl.rows[0].cells[0]
    sp_cell.text = ''
    p1 = sp_cell.paragraphs[0]
    run = p1.add_run('특기사항')
    run.font.bold = True
    run.font.size = Pt(9)
    run.font.name = '맑은 고딕'
    p2 = sp_cell.add_paragraph()
    run2 = p2.add_run(f"* 제품보존기간 : {d.get('qaPreservation','2 년')}")
    run2.font.size = Pt(9)
    run2.font.name = '맑은 고딕'
    
    doc.add_paragraph()
    
    dist_tbl = doc.add_table(rows=1, cols=2)
    dist_tbl.style = 'Table Grid'
    add_gray_cell(dist_tbl.rows[0].cells[0], '배 포 처', 8)
    cell_text(dist_tbl.rows[0].cells[1], d.get('qaDistribution',''), align='left', size=9)
    
    # 버퍼로 저장
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


class handler(BaseHTTPRequestHandler):
    def do_POST(self):
        length = int(self.headers.get('Content-Length', 0))
        body = self.rfile.read(length)
        data = json.loads(body)
        
        try:
            docx_bytes = generate_docx(data)
            b64 = base64.b64encode(docx_bytes).decode()
            
            self.send_response(200)
            self.send_header('Content-Type', 'application/json')
            self.send_header('Access-Control-Allow-Origin', '*')
            self.end_headers()
            self.wfile.write(json.dumps({'file': b64}).encode())
        except Exception as e:
            self.send_response(500)
            self.send_header('Content-Type', 'application/json')
            self.end_headers()
            self.wfile.write(json.dumps({'error': str(e)}).encode())
    
    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()
